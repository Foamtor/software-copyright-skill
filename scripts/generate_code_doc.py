#!/usr/bin/env python3
"""
生成代码文档脚本
提取项目代码，生成符合软著要求的代码文档
每页50行，前30页+后30页
"""

import os
import json
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 代码文件扩展名
CODE_EXTENSIONS = {
    '.java', '.py', '.js', '.ts', '.jsx', '.tsx', '.vue', '.go',
    '.rs', '.cpp', '.c', '.h', '.hpp', '.cs', '.php', '.rb',
    '.swift', '.kt', '.scala', '.lua', '.r', '.m', '.sql'
}

def create_code_document():
    """创建代码文档"""
    doc = Document()
    
    # 设置默认字体（代码用等宽字体）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(9)
    
    # 设置段落格式（单倍行距，无缩进）
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0
    paragraph_format.first_line_indent = Cm(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    return doc

def setup_header(doc, software_name, version):
    """设置页眉：左-软件名称，右-页码"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        # 清空现有页眉
        for para in header.paragraphs:
            para.clear()
        
        if not header.paragraphs:
            para = header.add_paragraph()
        else:
            para = header.paragraphs[0]
        
        para.alignment = 0  # 左对齐
        
        # 左侧：软件名称+版本号
        run = para.add_run(f"{software_name} {version}")
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 制表符
        tab_run = para.add_run('\t')
        
        # 右侧：页码
        add_page_field(para)

def add_page_field(paragraph):
    """插入页码域代码"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1 = paragraph.add_run()
    run1._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2 = paragraph.add_run()
    run2._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3 = paragraph.add_run()
    run3._element.append(fldChar2)

def add_file_header(doc, filename):
    """添加文件名标题"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加分隔线
    run = para.add_run('=' * 80)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # 添加文件名
    para = doc.add_paragraph()
    run = para.add_run(f'// 文件：{filename}')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.bold = True
    
    # 添加分隔线
    para = doc.add_paragraph()
    run = para.add_run('=' * 80)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

def add_code_content(doc, code_content, lines_per_page=50):
    """添加代码内容，控制每页行数"""
    lines = code_content.split('\n')
    
    for i, line in enumerate(lines):
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        
        # 每N行插入分页符
        if (i + 1) % lines_per_page == 0 and i < len(lines) - 1:
            doc.add_page_break()

def read_code_file(file_path, max_lines=None):
    """读取代码文件"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            if max_lines:
                lines = []
                for i, line in enumerate(f):
                    if i >= max_lines:
                        break
                    lines.append(line)
                return ''.join(lines)
            else:
                return f.read()
    except Exception as e:
        print(f"警告：无法读取文件 {file_path}: {e}")
        return None

def generate_code_doc(project_path, files_list, output_path, software_name, version, 
                      lines_per_page=50, max_pages=60):
    """
    生成代码文档
    
    Args:
        project_path: 项目路径
        files_list: 文件列表（JSON格式或列表）
        output_path: 输出文件路径
        software_name: 软件名称
        version: 版本号
        lines_per_page: 每页行数（默认50）
        max_pages: 最大页数（默认60）
    """
    doc = create_code_document()
    setup_header(doc, software_name, version)
    
    # 读取文件列表
    if isinstance(files_list, str):
        with open(files_list, 'r', encoding='utf-8') as f:
            files = json.load(f)
    else:
        files = files_list
    
    # 计算总行数限制
    total_lines_limit = lines_per_page * max_pages
    
    # 统计已添加的行数
    total_lines_added = 0
    
    for file_info in files:
        # 获取文件路径
        if isinstance(file_info, dict):
            file_path = file_info.get('path', '')
        else:
            file_path = file_info
        
        # 构建完整路径
        full_path = os.path.join(project_path, file_path)
        
        if not os.path.exists(full_path):
            print(f"警告：文件不存在 {full_path}")
            continue
        
        # 检查是否是代码文件
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in CODE_EXTENSIONS:
            continue
        
        # 读取代码
        code_content = read_code_file(full_path)
        if not code_content:
            continue
        
        # 计算当前文件行数
        file_lines = len(code_content.split('\n'))
        
        # 检查是否超过限制
        if total_lines_added + file_lines > total_lines_limit:
            # 截断文件
            remaining_lines = total_lines_limit - total_lines_added
            if remaining_lines > 0:
                lines = code_content.split('\n')[:remaining_lines]
                code_content = '\n'.join(lines)
                file_lines = remaining_lines
            else:
                break
        
        # 添加文件标题
        add_file_header(doc, file_path)
        
        # 添加代码内容
        add_code_content(doc, code_content, lines_per_page)
        
        total_lines_added += file_lines
        print(f"  已添加：{file_path} ({file_lines} 行)")
        
        # 检查是否达到限制
        if total_lines_added >= total_lines_limit:
            break
    
    # 保存文档
    doc.save(output_path)
    print(f"代码文档已保存到：{output_path}")
    print(f"  总行数：{total_lines_added}")

def main():
    parser = argparse.ArgumentParser(description='生成代码文档')
    parser.add_argument('--project', required=True, help='项目路径')
    parser.add_argument('--files', required=True, help='文件列表JSON路径')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--software-name', required=True, help='软件名称')
    parser.add_argument('--version', required=True, help='版本号')
    parser.add_argument('--lines-per-page', type=int, default=50, help='每页行数')
    parser.add_argument('--max-pages', type=int, default=60, help='最大页数')
    args = parser.parse_args()
    
    generate_code_doc(
        project_path=args.project,
        files_list=args.files,
        output_path=args.output,
        software_name=args.software_name,
        version=args.version,
        lines_per_page=args.lines_per_page,
        max_pages=args.max_pages
    )

if __name__ == '__main__':
    main()
