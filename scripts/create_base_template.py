#!/usr/bin/env python3
"""
创建基础Word模板（用于docxtpl）

这个脚本创建一个基础的Word模板，包含：
- 封面页（带Jinja2占位符）
- 正文样式定义
- 页眉页脚
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


def create_base_template(output_path):
    """创建基础模板"""
    doc = Document()
    
    # 1. 设置页面
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    
    # 2. 设置页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run("{{cover_title}}")
    run.font.name = "宋体"
    run.font.size = Pt(9)
    
    # 3. 创建样式
    # 正文内容样式
    if '正文内容' not in [s.name for s in doc.styles]:
        body_style = doc.styles.add_style('正文内容', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(16)
        body_style.paragraph_format.first_line_indent = Cm(1.18)
        body_style.paragraph_format.line_spacing = 1.5
        body_style.paragraph_format.space_before = Pt(0)
        body_style.paragraph_format.space_after = Pt(0)
    
    # 公司名称样式
    if '公司名称' not in [s.name for s in doc.styles]:
        company_style = doc.styles.add_style('公司名称', WD_STYLE_TYPE.PARAGRAPH)
        company_style.font.name = '宋体'
        company_style.font.size = Pt(16)
        company_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 4. 添加封面内容
    # 软件名称（大标题）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("{{cover_title}}")
    run.font.size = Pt(22)
    run.font.bold = True
    
    # 副标题
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("{{cover_subtitle}}")
    run.font.size = Pt(16)
    
    # 空行
    for _ in range(8):
        doc.add_paragraph()
    
    # 公司名称
    company_para = doc.add_paragraph(style='公司名称')
    run = company_para.add_run("{{cover_company}}")
    run.font.name = '宋体'
    run.font.size = Pt(16)
    
    # 日期
    date_para = doc.add_paragraph(style='公司名称')
    run = date_para.add_run("{{cover_date}}")
    run.font.name = '宋体'
    run.font.size = Pt(16)
    
    # 5. 分页
    doc.add_page_break()
    
    # 6. 添加章节占位（使用Jinja2循环）
    # 一级章节
    h1_para = doc.add_paragraph(style='Heading 1')
    run = h1_para.add_run("{% for section in sections %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)  # 白色，隐藏
    
    h1_para = doc.add_paragraph(style='Heading 1')
    run = h1_para.add_run("{{section.heading}}")
    run.font.name = '黑体'
    run.font.size = Pt(16)
    run.font.bold = False
    
    # 正文内容
    body_para = doc.add_paragraph(style='正文内容')
    run = body_para.add_run("{{section.content}}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    
    # 图片说明
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run("{% for img in section.images %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run("{{img.ref}} {{img.description}}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    img_para = doc.add_paragraph()
    run = img_para.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 注意事项
    note_para = doc.add_paragraph(style='正文内容')
    run = note_para.add_run("{% for note in section.notes %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    note_para = doc.add_paragraph(style='正文内容')
    run = note_para.add_run("注意：{{note}}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    
    note_para = doc.add_paragraph()
    run = note_para.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 子章节
    h2_para = doc.add_paragraph(style='Heading 2')
    run = h2_para.add_run("{% for sub in section.subsections %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    h2_para = doc.add_paragraph(style='Heading 2')
    run = h2_para.add_run("{{sub.heading}}")
    run.font.size = Pt(16)
    run.font.bold = True
    
    sub_body = doc.add_paragraph(style='正文内容')
    run = sub_body.add_run("{{sub.content}}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    
    sub_end = doc.add_paragraph()
    run = sub_end.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 章节结束
    section_end = doc.add_paragraph()
    run = section_end.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 保存
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"✅ 模板已创建: {output_path}")
    return output_path


if __name__ == "__main__":
    import sys
    output = sys.argv[1] if len(sys.argv) > 1 else "templates/manual_template.docx"
    create_base_template(output)
