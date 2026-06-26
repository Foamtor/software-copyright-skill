#!/usr/bin/env python3
"""
公共工具函数库

提供各个脚本共享的功能：
1. 文档操作（加载、保存、段落处理）
2. 格式应用（字体、段落格式）
3. 图片处理（插入、映射）
4. 配置管理（加载、验证）
"""

import json
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少依赖：pip install python-docx")
    exit(1)


# ============================================================================
# 配置管理
# ============================================================================

def load_config(config_path):
    """加载格式规范配置"""
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_content(content_path):
    """加载JSON内容"""
    with open(content_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def validate_content(content):
    """验证JSON内容格式"""
    errors = []
    
    if not content.get("sections"):
        errors.append("缺少章节内容 (sections)")
    
    # 检查封面信息
    if not content.get("cover", {}).get("title"):
        if not content.get("software_name"):
            errors.append("缺少封面标题 (cover.title 或 software_name)")
    
    for i, section in enumerate(content.get("sections", [])):
        if not section.get("heading"):
            errors.append(f"章节 {i+1} 缺少标题 (heading)")
        
        if not section.get("content") and not section.get("subsections"):
            errors.append(f"章节 {i+1} 缺少内容 (content 或 subsections)")
    
    return errors


# ============================================================================
# 文档操作
# ============================================================================

def load_document(doc_path):
    """加载Word文档"""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"文档不存在: {doc_path}")
    return Document(doc_path)


def save_document(doc, output_path):
    """保存Word文档"""
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    return output_path


def get_paragraph_text(para):
    """获取段落文本（去除首尾空格）"""
    return para.text.strip()


def is_paragraph_empty(para):
    """判断段落是否为空"""
    return not para.text.strip()


def find_paragraphs_by_style(doc, style_name):
    """按样式查找段落"""
    return [p for p in doc.paragraphs if p.style.name == style_name]


def find_paragraph_by_text(doc, text_pattern):
    """按文本内容查找段落"""
    for i, para in enumerate(doc.paragraphs):
        if re.search(text_pattern, para.text):
            return i, para
    return -1, None


# ============================================================================
# 段落操作
# ============================================================================

def insert_paragraph_before(paragraph):
    """在指定段落前插入新段落"""
    new_p = OxmlElement('w:p')
    paragraph._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after(paragraph):
    """在指定段落后插入新段落"""
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def copy_paragraph_format(source_para, target_para):
    """复制段落格式"""
    target_para.style = source_para.style
    target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
    target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
    target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    target_para.paragraph_format.space_after = source_para.paragraph_format.space_after


def copy_run_format(source_run, target_run):
    """复制Run格式"""
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.bold = source_run.font.bold
    target_run.font.color.rgb = source_run.font.color.rgb
    
    # 设置中文字体
    r = target_run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')


# ============================================================================
# 格式应用
# ============================================================================

def set_font(run, font_name, font_size_pt, bold=False, color="000000"):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_alignment(para, alignment):
    """设置段落对齐方式"""
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    if isinstance(alignment, str):
        para.paragraph_format.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    else:
        para.paragraph_format.alignment = alignment


def set_paragraph_indent(para, first_line_indent_cm=None, left_indent_cm=None, right_indent_cm=None):
    """设置段落缩进"""
    if first_line_indent_cm is not None:
        para.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if left_indent_cm is not None:
        para.paragraph_format.left_indent = Cm(left_indent_cm)
    if right_indent_cm is not None:
        para.paragraph_format.right_indent = Cm(right_indent_cm)


def set_paragraph_spacing(para, line_spacing=None, space_before_pt=None, space_after_pt=None):
    """设置段落间距"""
    if line_spacing is not None:
        para.paragraph_format.line_spacing = line_spacing
    if space_before_pt is not None:
        para.paragraph_format.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        para.paragraph_format.space_after = Pt(space_after_pt)


# ============================================================================
# 页眉页脚
# ============================================================================

def add_page_number(paragraph):
    """在段落中添加页码域代码"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)


def replace_header_footer(doc, software_name, version):
    """替换页眉页脚占位符，添加页码"""
    for section in doc.sections:
        # 页眉：替换占位符
        for para in section.header.paragraphs:
            for run in para.runs:
                if '{software_name}' in run.text:
                    run.text = run.text.replace('{software_name}', software_name)
                if '{version}' in run.text:
                    run.text = run.text.replace('{version}', version)
        
        # 页脚：添加页码（居中）
        for para in section.footer.paragraphs:
            # 清空现有内容
            for run in para.runs:
                run.text = ''
            # 添加页码
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_page_number(para)


# ============================================================================
# 图片处理
# ============================================================================

def build_image_mapping(content, image_dir):
    """构建图片映射：图号 -> 图片路径"""
    image_mapping = {}
    
    for section in content.get('sections', []):
        for img in section.get('images', []):
            ref = img.get('ref', '')
            path = img.get('path', '')
            if ref and path:
                full_path = os.path.join(image_dir, path)
                if os.path.exists(full_path):
                    image_mapping[ref] = full_path
        
        for sub in section.get('subsections', []):
            for img in sub.get('images', []):
                ref = img.get('ref', '')
                path = img.get('path', '')
                if ref and path:
                    full_path = os.path.join(image_dir, path)
                    if os.path.exists(full_path):
                        image_mapping[ref] = full_path
    
    return image_mapping


def insert_image_before_paragraph(para, img_path, width_inches=5.0):
    """在段落前插入图片"""
    new_para = insert_paragraph_before(para)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return new_para


def insert_image_after_paragraph(para, img_path, width_inches=5.0):
    """在段落后插入图片"""
    new_para = insert_paragraph_after(para)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return new_para


# ============================================================================
# 段落分割
# ============================================================================

def split_long_paragraphs(doc):
    """将长段落分割成多个独立段落"""
    paragraphs = list(doc.paragraphs)
    split_count = 0
    
    for i in range(len(paragraphs)-1, -1, -1):
        para = paragraphs[i]
        text = para.text
        
        if '\n' in text:
            lines = text.split('\n')
            main_text = lines[0]
            other_lines = [line for line in lines[1:] if line.strip()]
            
            if other_lines:
                # 清空原段落
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = main_text
                else:
                    para.add_run(main_text)
                
                # 在原段落后插入新段落
                prev_para = para
                for line in other_lines:
                    new_para = insert_paragraph_after(prev_para)
                    
                    # 复制原段落的格式
                    copy_paragraph_format(para, new_para)
                    
                    # 设置文本
                    run = new_para.add_run(line)
                    
                    # 复制字体格式
                    if para.runs:
                        copy_run_format(para.runs[0], run)
                    
                    prev_para = new_para
                    split_count += 1
    
    return split_count


def split_paragraph_at_figure_captions(doc):
    """将图名分割成独立段落"""
    paragraphs = list(doc.paragraphs)
    split_count = 0
    
    for i in range(len(paragraphs)-1, -1, -1):
        para = paragraphs[i]
        text = para.text
        
        # 查找图名
        fig_pattern = r'(图\d+\s+[^\n]+)'
        matches = list(re.finditer(fig_pattern, text))
        
        if len(matches) > 0:
            last_end = len(text)
            fig_captions = []
            
            for match in reversed(matches):
                fig_captions.insert(0, match.group(1))
                last_end = match.start()
            
            main_text = text[:last_end].strip()
            
            if main_text:
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = main_text
                else:
                    para.add_run(main_text)
                
                prev_para = para
                for fig_caption in fig_captions:
                    new_para = insert_paragraph_after(prev_para)
                    
                    # 设置图名格式：14pt，居中
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = new_para.add_run(fig_caption)
                    run.font.size = Pt(14)
                    
                    prev_para = new_para
                    split_count += 1
    
    return split_count


def remove_consecutive_empty_lines(doc):
    """删除连续空行，只保留一个"""
    paragraphs = list(doc.paragraphs)
    to_remove = []
    
    for i in range(len(paragraphs) - 1):
        curr_empty = is_paragraph_empty(paragraphs[i])
        next_empty = is_paragraph_empty(paragraphs[i+1])
        
        if curr_empty and next_empty:
            to_remove.append(i+1)
    
    # 从后往前删除
    for i in reversed(to_remove):
        p = paragraphs[i]._element
        p.getparent().remove(p)
    
    return len(to_remove)


# ============================================================================
# 内容处理
# ============================================================================

def count_chars(text):
    """统计字数（不含空格和换行）"""
    return len(text.replace(" ", "").replace("\n", ""))


def count_section_chars(section):
    """统计章节字数"""
    total = count_chars(section.get('content', ''))
    
    for sub in section.get('subsections', []):
        total += count_chars(sub.get('content', ''))
    
    return total


def extract_figure_refs(text):
    """提取文本中的图号"""
    return re.findall(r'图\d+', text)


def has_figure_caption(text):
    """检查是否包含图名"""
    return bool(re.search(r'图\d+\s+', text))


# ============================================================================
# 文件操作
# ============================================================================

def ensure_dir(dir_path):
    """确保目录存在"""
    os.makedirs(dir_path, exist_ok=True)
    return dir_path


def get_file_size(file_path):
    """获取文件大小（KB）"""
    return os.path.getsize(file_path) / 1024


def list_files(dir_path, pattern="*"):
    """列出目录下的文件"""
    return list(Path(dir_path).glob(pattern))


# ============================================================================
# Brand Profile 支持
# ============================================================================

def load_brand_profile(profile_path):
    """加载Brand Profile配置"""
    with open(profile_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_style_config(brand_profile, style_name):
    """从Brand Profile获取样式配置"""
    styles = brand_profile.get('styles', {})
    return styles.get(style_name, {})


def get_font_config(style_config):
    """从样式配置获取字体配置"""
    return style_config.get('font', {})


def get_paragraph_config(style_config):
    """从样式配置获取段落配置"""
    return style_config.get('paragraph', {})


def apply_brand_style(para, style_config):
    """应用Brand Profile样式到段落"""
    font_config = get_font_config(style_config)
    para_config = get_paragraph_config(style_config)
    
    # 设置字体
    if font_config:
        for run in para.runs:
            if font_config.get('name'):
                run.font.name = font_config['name']
            if font_config.get('size_pt'):
                run.font.size = Pt(font_config['size_pt'])
            if font_config.get('bold') is not None:
                run.font.bold = font_config['bold']
            if font_config.get('color'):
                run.font.color.rgb = RGBColor.from_string(font_config['color'].lstrip('#'))
            
            # 设置中文字体
            east_asia = font_config.get('east_asia') or font_config.get('name_cn')
            if east_asia:
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), east_asia)
    
    # 设置段落格式
    if para_config:
        if para_config.get('alignment'):
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            para.paragraph_format.alignment = align_map.get(para_config['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
        
        if para_config.get('first_line_indent_cm') is not None:
            para.paragraph_format.first_line_indent = Cm(para_config['first_line_indent_cm'])
        
        if para_config.get('line_spacing') is not None:
            para.paragraph_format.line_spacing = para_config['line_spacing']
        
        if para_config.get('space_before_pt') is not None:
            para.paragraph_format.space_before = Pt(para_config['space_before_pt'])
        
        if para_config.get('space_after_pt') is not None:
            para.paragraph_format.space_after = Pt(para_config['space_after_pt'])


def create_styled_paragraph(doc, text, style_config, style_name=None):
    """创建带样式的段落"""
    if style_name:
        para = doc.add_paragraph(style=style_name)
    else:
        para = doc.add_paragraph()
    
    run = para.add_run(text)
    apply_brand_style(para, style_config)
    
    return para


def brand_profile_to_format_spec(brand_profile):
    """将Brand Profile转换为旧的format_spec格式（兼容）"""
    styles = brand_profile.get('styles', {})
    
    format_spec = {
        'version': brand_profile.get('version', '1.0'),
        'page': brand_profile.get('page_layout', {}),
        'header': styles.get('header', {}),
        'cover': {
            'title': styles.get('cover_title', {}),
            'subtitle': styles.get('cover_subtitle', {}),
            'company': styles.get('cover_company', {}),
            'date': styles.get('cover_date', {})
        },
        'heading1': styles.get('heading1', {}),
        'heading2': styles.get('heading2', {}),
        'heading3': styles.get('heading3', {}),
        'body': styles.get('body', {}),
        'figure_caption': styles.get('figure_caption', {}),
        'emu_conversions': brand_profile.get('emu_conversions', {})
    }
    
    return format_spec
