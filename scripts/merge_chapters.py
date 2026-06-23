#!/usr/bin/env python3
"""
合并章节文档脚本
将多个章节文档合并为完整的说明书
"""

import os
import json
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_header(doc, software_name, version):
    """设置页眉：左-软件名称，右-页码"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        # 清空现有页眉
        for para in header.paragraphs:
            para.clear()
        
        if not header.paragraphs:
            para = header.add_paragraph()
        else:
            para = header.paragraphs[0]
        
        para.alignment = 0  # 左对齐
        
        # 左侧：软件名称+版本号
        run = para.add_run(f"{software_name} {version}")
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 制表符
        tab_run = para.add_run('\t')
        
        # 右侧：页码
        add_page_field(para)

def add_page_field(paragraph):
    """插入页码域代码"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1 = paragraph.add_run()
    run1._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2 = paragraph.add_run()
    run2._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3 = paragraph.add_run()
    run3._element.append(fldChar2)

def copy_paragraph(source_para, target_doc):
    """复制段落到目标文档"""
    # 创建新段落
    new_para = target_doc.add_paragraph()
    
    # 复制段落格式
    if source_para.style:
        new_para.style = source_para.style
    
    # 复制对齐方式
    new_para.alignment = source_para.alignment
    
    # 复制段落格式属性
    if source_para.paragraph_format:
        new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    
    # 复制runs
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        if run.bold:
            new_run.bold = run.bold
        if run.italic:
            new_run.italic = run.italic
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.size:
            new_run.font.size = run.font.size
    
    return new_para

def copy_table(source_table, target_doc):
    """复制表格到目标文档"""
    # 创建新表格
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    new_table = target_doc.add_table(rows=rows, cols=cols)
    
    # 复制表格样式
    if source_table.style:
        new_table.style = source_table.style
    
    # 复制单元格内容
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            # 清空新单元格
            for para in new_cell.paragraphs:
                para.clear()
            # 复制内容
            for k, para in enumerate(cell.paragraphs):
                if k == 0:
                    new_para = new_cell.paragraphs[0]
                else:
                    new_para = new_cell.add_paragraph()
                
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size
    
    return new_table

def merge_documents(chapter_files, output_path, software_name=None, version=None):
    """
    合并多个章节文档
    
    Args:
        chapter_files: 章节文件列表
        output_path: 输出文件路径
        software_name: 软件名称（用于页眉）
        version: 版本号（用于页眉）
    """
    # 创建目标文档
    merged_doc = Document()
    
    # 设置默认字体
    style = merged_doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 复制每个章节
    for i, chapter_file in enumerate(chapter_files):
        if not os.path.exists(chapter_file):
            print(f"警告：文件不存在 {chapter_file}")
            continue
        
        print(f"  合并：{chapter_file}")
        
        # 加载章节文档
        chapter_doc = Document(chapter_file)
        
        # 复制段落
        for para in chapter_doc.paragraphs:
            copy_paragraph(para, merged_doc)
        
        # 复制表格
        for table in chapter_doc.tables:
            copy_table(table, merged_doc)
        
        # 添加分页符（除了最后一章）
        if i < len(chapter_files) - 1:
            merged_doc.add_page_break()
    
    # 设置页眉
    if software_name and version:
        setup_header(merged_doc, software_name, version)
    
    # 保存合并后的文档
    merged_doc.save(output_path)
    print(f"合并完成：{output_path}")

def main():
    parser = argparse.ArgumentParser(description='合并章节文档')
    parser.add_argument('--chapters', nargs='+', required=True, help='章节文件列表')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--software-name', help='软件名称（用于页眉）')
    parser.add_argument('--version', help='版本号（用于页眉）')
    parser.add_argument('--chapter-list', help='章节文件列表JSON')
    args = parser.parse_args()
    
    # 获取章节文件列表
    chapter_files = args.chapters
    
    # 如果提供了JSON列表
    if args.chapter_list:
        with open(args.chapter_list, 'r', encoding='utf-8') as f:
            chapter_files = json.load(f)
    
    merge_documents(
        chapter_files=chapter_files,
        output_path=args.output,
        software_name=args.software_name,
        version=args.version
    )

if __name__ == '__main__':
    main()
