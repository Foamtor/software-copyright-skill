#!/usr/bin/env python3
"""
模板填充脚本
读取Word模板，替换占位符，保持原有格式
"""

import os
import json
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def replace_in_paragraph(paragraph, placeholder, value):
    """在段落中替换占位符（run级别，保持格式）"""
    if placeholder not in paragraph.text:
        return False
    
    # 遍历所有run
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, str(value))
            return True
    
    # 如果占位符跨多个run，需要合并处理
    # 这种情况较少见，暂时跳过
    return False

def replace_in_table(table, placeholder, value):
    """在表格中替换占位符"""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, placeholder, value):
                    replaced = True
    return replaced

def replace_all(doc, placeholder, value):
    """在整个文档中替换占位符"""
    replaced = False
    
    # 替换正文中的占位符
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, placeholder, value):
            replaced = True
    
    # 替换表格中的占位符
    for table in doc.tables:
        if replace_in_table(table, placeholder, value):
            replaced = True
    
    # 替换页眉中的占位符
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if replace_in_paragraph(paragraph, placeholder, value):
                replaced = True
        
        footer = section.footer
        for paragraph in footer.paragraphs:
            if replace_in_paragraph(paragraph, placeholder, value):
                replaced = True
    
    return replaced

def setup_header(doc, software_name, version):
    """设置页眉：左-软件名称，右-页码"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        # 清空现有页眉内容
        for para in header.paragraphs:
            para.clear()
        
        # 如果没有段落，添加一个
        if not header.paragraphs:
            para = header.add_paragraph()
        else:
            para = header.paragraphs[0]
        
        # 设置对齐
        para.alignment = 0  # 左对齐
        
        # 左侧：软件名称+版本号
        run = para.add_run(f"{software_name} {version}")
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加制表符（右对齐）
        tab_run = para.add_run('\t')
        
        # 右侧：页码域代码
        add_page_field(para)

def add_page_field(paragraph):
    """插入页码域代码"""
    # 开始域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1 = paragraph.add_run()
    run1._element.append(fldChar1)
    
    # 域代码
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2 = paragraph.add_run()
    run2._element.append(instrText)
    
    # 结束域
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3 = paragraph.add_run()
    run3._element.append(fldChar2)

def fill_template(template_path, content, output_path, setup_hdr=True, software_name=None, version=None):
    """
    填充模板
    
    Args:
        template_path: 模板文件路径
        content: 内容字典 {占位符: 值}
        output_path: 输出文件路径
        setup_hdr: 是否设置页眉
        software_name: 软件名称（用于页眉）
        version: 版本号（用于页眉）
    """
    # 加载模板
    doc = Document(template_path)
    
    # 替换所有占位符
    for placeholder, value in content.items():
        # 确保占位符格式为 {{xxx}}
        if not placeholder.startswith('{{'):
            placeholder = '{{' + placeholder + '}}'
        
        replace_all(doc, placeholder, value)
    
    # 设置页眉
    if setup_hdr and software_name and version:
        setup_header(doc, software_name, version)
    
    # 保存文档
    doc.save(output_path)
    print(f"文档已保存到：{output_path}")

def main():
    parser = argparse.ArgumentParser(description='填充Word模板')
    parser.add_argument('--template', required=True, help='模板文件路径')
    parser.add_argument('--content', required=True, help='内容JSON文件路径')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--software-name', help='软件名称（用于页眉）')
    parser.add_argument('--version', help='版本号（用于页眉）')
    parser.add_argument('--no-header', action='store_true', help='不设置页眉')
    args = parser.parse_args()
    
    # 读取内容
    with open(args.content, 'r', encoding='utf-8') as f:
        content = json.load(f)
    
    # 填充模板
    fill_template(
        template_path=args.template,
        content=content,
        output_path=args.output,
        setup_hdr=not args.no_header,
        software_name=args.software_name,
        version=args.version
    )

if __name__ == '__main__':
    main()
