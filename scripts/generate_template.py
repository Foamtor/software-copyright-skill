#!/usr/bin/env python3
"""
生成自定义模板脚本
分析用户提供的参考文档，生成新的Word模板
"""

import os
import json
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import re

def analyze_document_structure(doc_path):
    """分析Word文档结构"""
    doc = Document(doc_path)
    
    result = {
        'paragraphs': [],
        'tables': [],
        'styles': set(),
        'fonts': set()
    }
    
    # 分析段落
    for i, para in enumerate(doc.paragraphs):
        para_info = {
            'index': i,
            'text': para.text[:100] if para.text else '',
            'style': para.style.name if para.style else 'Normal',
            'runs': []
        }
        
        for run in para.runs:
            run_info = {
                'text': run.text[:50] if run.text else '',
                'font': run.font.name,
                'size': str(run.font.size) if run.font.size else None,
                'bold': run.font.bold
            }
            para_info['runs'].append(run_info)
            
            if run.font.name:
                result['fonts'].add(run.font.name)
        
        result['paragraphs'].append(para_info)
        result['styles'].add(para_info['style'])
    
    # 分析表格
    for t_idx, table in enumerate(doc.tables):
        table_info = {
            'index': t_idx,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'cells': []
        }
        
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_info = {
                    'row': r_idx,
                    'col': c_idx,
                    'text': cell.text[:50] if cell.text else ''
                }
                table_info['cells'].append(cell_info)
        
        result['tables'].append(table_info)
    
    result['styles'] = list(result['styles'])
    result['fonts'] = list(result['fonts'])
    
    return result

def find_placeholders(text):
    """查找文本中的占位符"""
    # 匹配 {{xxx}} 格式
    pattern = r'\{\{([^}]+)\}\}'
    return re.findall(pattern, text)

def create_template_from_reference(ref_doc_path, output_path, placeholder_mapping=None):
    """
    从参考文档创建模板
    
    Args:
        ref_doc_path: 参考文档路径
        output_path: 输出模板路径
        placeholder_mapping: 占位符映射 {原内容: 占位符名称}
    """
    # 加载参考文档
    doc = Document(ref_doc_path)
    
    # 分析结构
    structure = analyze_document_structure(ref_doc_path)
    
    # 替换内容为占位符
    if placeholder_mapping:
        # 替换段落中的内容
        for para in doc.paragraphs:
            for run in para.runs:
                for original, placeholder in placeholder_mapping.items():
                    if original in run.text:
                        run.text = run.text.replace(original, f'{{{{{placeholder}}}}}')
        
        # 替换表格中的内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for original, placeholder in placeholder_mapping.items():
                                if original in run.text:
                                    run.text = run.text.replace(original, f'{{{{{placeholder}}}}}')
    
    # 保存模板
    doc.save(output_path)
    print(f"模板已保存到：{output_path}")
    
    return structure

def generate_template_from_scratch(template_type, output_path):
    """
    从零生成模板
    
    Args:
        template_type: 模板类型（'collection_form', 'manual', 'code'）
        output_path: 输出路径
    """
    doc = Document()
    
    if template_type == 'collection_form':
        # 生成信息采集表模板
        create_collection_form_template(doc)
    elif template_type == 'manual':
        # 生成说明书模板
        create_manual_template(doc)
    elif template_type == 'code':
        # 生成代码文档模板
        create_code_template(doc)
    
    doc.save(output_path)
    print(f"模板已保存到：{output_path}")

def create_collection_form_template(doc):
    """创建信息采集表模板"""
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    title = doc.add_heading('软件著作权登记申请表', 0)
    for run in title.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 表1：软件基本信息
    doc.add_heading('软件基本信息', level=1)
    
    table1 = doc.add_table(rows=12, cols=4)
    table1.style = 'Table Grid'
    
    # 填充表头
    headers = ['序号', '项目', '填写内容', '说明']
    for i, header in enumerate(headers):
        cell = table1.cell(0, i)
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    # 填充字段
    fields = [
        ('1', '软件全称', '{{软件全称}}', '须以"系统、软件、平台"结尾'),
        ('2', '软件简称', '{{软件简称}}', '字数少于全称'),
        ('3', '版本号', '{{版本号}}', '如 V1.0'),
        ('4', '软件分类', '{{软件分类}}', '应用软件/嵌入式软件/中间件/操作系统'),
        ('5', '软件作品说明', '{{软件作品说明}}', '原创/修改'),
        ('6', '开发方式', '{{开发方式}}', '独立开发/合作开发/委托开发'),
        ('7', '著作权人', '{{著作权人}}', '姓名或名称'),
        ('8', '开发完成日期', '{{开发完成日期}}', '不得早于著作权人成立日期'),
        ('9', '发表状态', '{{发表状态}}', '未发表/已发表'),
        ('10', '已发表信息', '{{已发表信息}}', '首次发表日期、地点'),
        ('11', '联系人信息', '{{联系人信息}}', '姓名、邮箱、电话、地址'),
    ]
    
    for i, (num, field, placeholder, desc) in enumerate(fields):
        row = table1.rows[i + 1]
        row.cells[0].text = num
        row.cells[1].text = field
        row.cells[2].text = placeholder
        row.cells[3].text = desc
    
    # 表2：开发与运行环境
    doc.add_heading('开发与运行环境', level=1)
    
    table2 = doc.add_table(rows=10, cols=2)
    table2.style = 'Table Grid'
    
    # 填充字段
    env_fields = [
        ('开发的硬件环境', '{{开发硬件环境}}'),
        ('运行的硬件环境', '{{运行硬件环境}}'),
        ('开发操作系统', '{{开发操作系统}}'),
        ('开发环境/工具', '{{开发工具}}'),
        ('运行平台/操作系统', '{{运行平台}}'),
        ('运行支撑环境', '{{运行支撑环境}}'),
        ('编程语言', '{{编程语言}}'),
        ('源程序量', '{{源程序量}}'),
        ('开发目的', '{{开发目的}}'),
    ]
    
    for i, (field, placeholder) in enumerate(env_fields):
        row = table2.rows[i]
        row.cells[0].text = field
        row.cells[1].text = placeholder

def create_manual_template(doc):
    """创建说明书模板"""
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    title = doc.add_heading('{{软件名称}} {{版本号}} 使用说明书', 0)
    for run in title.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 添加示例章节结构
    doc.add_heading('一、系统概述', level=1)
    doc.add_heading('（一）系统简介', level=2)
    doc.add_heading('1. 系统定位', level=3)
    doc.add_paragraph('{{系统定位内容}}')
    doc.add_heading('2. 目标用户', level=3)
    doc.add_paragraph('{{目标用户内容}}')
    
    doc.add_heading('（二）功能概览', level=2)
    doc.add_paragraph('{{功能概览内容}}')
    
    doc.add_heading('二、用户端功能', level=1)
    doc.add_heading('（一）用户登录', level=2)
    doc.add_heading('1. 登录入口', level=3)
    doc.add_paragraph('{{登录入口内容}}')
    doc.add_heading('2. 登录流程', level=3)
    doc.add_paragraph('{{登录流程内容}}')
    
    # 添加截图占位符示例
    doc.add_paragraph('{{登录页面截图}}')
    doc.add_paragraph('{{截图说明}}')

def create_code_template(doc):
    """创建代码文档模板"""
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(9)
    
    # 添加标题
    title = doc.add_heading('{{软件名称}} {{版本号}} 源代码文档', 0)
    for run in title.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 添加代码示例
    doc.add_paragraph('// 文件：{{文件名}}')
    doc.add_paragraph('{{代码内容}}')

def main():
    parser = argparse.ArgumentParser(description='生成自定义模板')
    parser.add_argument('--type', choices=['collection_form', 'manual', 'code', 'from_ref'], 
                       default='from_ref', help='模板类型')
    parser.add_argument('--ref-doc', help='参考文档路径')
    parser.add_argument('--output', required=True, help='输出模板路径')
    parser.add_argument('--mapping', help='占位符映射JSON文件')
    args = parser.parse_args()
    
    if args.type == 'from_ref':
        if not args.ref_doc:
            print("错误：从参考文档生成需要提供 --ref-doc 参数")
            return
        
        # 读取映射
        mapping = None
        if args.mapping:
            with open(args.mapping, 'r', encoding='utf-8') as f:
                mapping = json.load(f)
        
        create_template_from_reference(args.ref_doc, args.output, mapping)
    else:
        generate_template_from_scratch(args.type, args.output)

if __name__ == '__main__':
    main()
