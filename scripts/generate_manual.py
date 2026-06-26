#!/usr/bin/env python3
"""
JSON + 模板 → Word文档生成器（完整版）

用法：
  python3 generate_manual.py \
    --content content.json \
    --template templates/manual_template.docx \
    --output output/说明书.docx \
    [--config config/brand_profile.json] \
    [--image-dir images/]

功能：
  1. 读取JSON内容文件
  2. 使用docxtpl渲染Word模板
  3. 替换页眉页脚占位符，添加页码
  4. 在图名上方插入实际图片
  5. 分割长段落，确保首行缩进
  6. 删除连续空行
  7. 使用Brand Profile应用格式（跳过封面）
  8. 输出格式规范的Word文档
"""

import argparse
import json
import os
import re
from pathlib import Path

try:
    from docxtpl import DocxTemplate
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少依赖：pip install docxtpl python-docx")
    exit(1)

# 导入公共工具
from utils import (
    load_config, load_content, validate_content, load_brand_profile,
    load_document, save_document,
    insert_paragraph_before, insert_paragraph_after,
    copy_paragraph_format, copy_run_format,
    set_font, set_paragraph_alignment, set_paragraph_indent, set_paragraph_spacing,
    add_page_number, replace_header_footer,
    build_image_mapping, insert_image_before_paragraph,
    split_long_paragraphs, split_paragraph_at_figure_captions, remove_consecutive_empty_lines,
    count_chars, get_style_config, apply_brand_style, brand_profile_to_format_spec
)


class ManualGenerator:
    """说明书生成器"""
    
    def __init__(self, template_path, config_path=None):
        """初始化生成器"""
        self.template_path = template_path
        self.doc = DocxTemplate(template_path)
        
        # 加载配置（支持Brand Profile和旧格式）
        self.config = None
        self.brand_profile = None
        
        if config_path and os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                config_data = json.load(f)
            
            # 检测是否是Brand Profile格式
            if 'styles' in config_data and 'brand_info' in config_data:
                self.brand_profile = config_data
                self.config = brand_profile_to_format_spec(config_data)
            else:
                self.config = config_data
    
    def prepare_context(self, content):
        """准备模板上下文"""
        # 封面信息
        cover = content.get("cover", {})
        software_name = content.get("software_name", "")
        version = content.get("version", "")
        company = content.get("company", "")
        
        context = {
            "cover_title": cover.get("title", f"{software_name}{version}"),
            "cover_subtitle": cover.get("subtitle", "用户使用手册"),
            "cover_company": cover.get("company", company),
            "cover_date": cover.get("date", ""),
            "sections": []
        }
        
        for section in content.get("sections", []):
            section_data = {
                "heading": section.get("heading", ""),
                "level": section.get("level", 1),
                "content": section.get("content", ""),
                "images": section.get("images", []),
                "notes": section.get("notes", []),
                "subsections": []
            }
            
            for sub in section.get("subsections", []):
                sub_data = {
                    "heading": sub.get("heading", ""),
                    "content": sub.get("content", ""),
                    "images": sub.get("images", []),
                    "notes": sub.get("notes", [])
                }
                section_data["subsections"].append(sub_data)
            
            context["sections"].append(section_data)
        
        return context
    
    def apply_formatting(self, doc):
        """应用格式到文档（跳过封面段落）"""
        if not self.config:
            return
        
        # 找到第一个Heading 1的位置，之前的都是封面
        first_heading_idx = 0
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == 'Heading 1':
                first_heading_idx = i
                break
        
        # 遍历所有段落，应用格式
        for para_idx, para in enumerate(doc.paragraphs):
            style_name = para.style.name
            text = para.text.strip()
            
            if not text:
                continue
            
            # 跳过封面段落
            if para_idx < first_heading_idx:
                continue
            
            # 根据样式类型应用格式
            if style_name == 'Heading 1':
                style_config = get_style_config(self.brand_profile, 'heading1') if self.brand_profile else self.config.get('heading1', {})
                apply_brand_style(para, style_config)
            elif style_name == 'Heading 2':
                style_config = get_style_config(self.brand_profile, 'heading2') if self.brand_profile else self.config.get('heading2', {})
                apply_brand_style(para, style_config)
            elif text.startswith('图') and any(c.isdigit() for c in text[:5]):
                style_config = get_style_config(self.brand_profile, 'figure_caption') if self.brand_profile else self.config.get('figure_caption', {})
                apply_brand_style(para, style_config)
            else:
                style_config = get_style_config(self.brand_profile, 'body') if self.brand_profile else self.config.get('body', {})
                apply_brand_style(para, style_config)
    
    def generate(self, content, output_path, image_dir=None):
        """生成Word文档"""
        # 准备上下文
        context = self.prepare_context(content)
        
        # 渲染模板
        self.doc.render(context)
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        
        # 保存文档
        self.doc.save(output_path)
        
        # 加载文档进行后处理
        print("\n📝 后处理...")
        doc = Document(output_path)
        
        # 替换页眉页脚
        software_name = content.get("software_name", "")
        version = content.get("version", "")
        replace_header_footer(doc, software_name, version)
        
        # 1. 分割长段落
        split_count = split_long_paragraphs(doc)
        print(f"   分割长段落: {split_count}")
        
        # 2. 分割图名
        fig_count = split_paragraph_at_figure_captions(doc)
        print(f"   分割图名: {fig_count}")
        
        # 3. 插入图片
        if image_dir:
            image_mapping = build_image_mapping(content, image_dir)
            insert_count = 0
            paragraphs = list(doc.paragraphs)
            inserted_refs = set()
            
            for i in range(len(paragraphs)-1, -1, -1):
                para = paragraphs[i]
                text = para.text.strip()
                
                match = re.match(r'^(图\d+)\s+', text)
                if match:
                    fig_ref = match.group(1)
                    if fig_ref in image_mapping and fig_ref not in inserted_refs:
                        img_path = image_mapping[fig_ref]
                        insert_image_before_paragraph(para, img_path)
                        insert_count += 1
                        inserted_refs.add(fig_ref)
            
            print(f"   插入图片: {insert_count}")
        
        # 4. 删除连续空行
        empty_count = remove_consecutive_empty_lines(doc)
        print(f"   删除空行: {empty_count}")
        
        # 5. 应用格式
        if self.config:
            self.apply_formatting(doc)
        
        # 保存文档
        doc.save(output_path)
        
        print(f"\n✅ 文档已生成: {output_path}")
        print(f"   模板: {self.template_path}")
        print(f"   章节数: {len(context['sections'])}")
        
        return output_path


def main():
    parser = argparse.ArgumentParser(description="JSON + 模板 → Word文档")
    parser.add_argument("--content", required=True, help="内容JSON文件路径")
    parser.add_argument("--template", required=True, help="Word模板文件路径")
    parser.add_argument("--output", required=True, help="输出Word文件路径")
    parser.add_argument("--config", default="config/brand_profile.json", help="Brand Profile或format_spec配置文件路径")
    parser.add_argument("--image-dir", help="图片目录路径")
    parser.add_argument("--validate", action="store_true", help="仅验证，不生成")
    
    args = parser.parse_args()
    
    # 检查文件是否存在
    if not os.path.exists(args.content):
        print(f"❌ 内容文件不存在: {args.content}")
        exit(1)
    
    if not os.path.exists(args.template):
        print(f"❌ 模板文件不存在: {args.template}")
        exit(1)
    
    # 加载内容
    content = load_content(args.content)
    
    if args.validate:
        errors = validate_content(content)
        if errors:
            print("⚠️ 内容格式警告：")
            for error in errors:
                print(f"   - {error}")
        else:
            print("✅ 内容格式验证通过")
        exit(0)
    
    # 生成文档
    generator = ManualGenerator(args.template, args.config)
    generator.generate(content, args.output, args.image_dir)


if __name__ == "__main__":
    main()
