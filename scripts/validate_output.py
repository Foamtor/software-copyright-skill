#!/usr/bin/env python3
"""
软著文档格式验证脚本（使用配置文件）。

用法：
  python3 validate_output.py --docx output/说明书.docx [--config config/format_spec.json]
  python3 validate_output.py --docx output/说明书.docx --fix [--config config/format_spec.json]
  python3 validate_output.py --docx output/说明书.docx --report report.json

功能：
  1. 读取格式规范配置文件
  2. 检查文档格式是否符合规范
  3. 可选自动修复格式
"""

import argparse
import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少依赖：pip install python-docx")
    exit(1)


def load_config(config_path):
    """加载格式规范配置"""
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)


class FormatValidator:
    """格式验证器"""
    
    def __init__(self, docx_path, config_path=None):
        """初始化"""
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.errors = []
        self.warnings = []
        self.stats = {
            "paragraphs": 0,
            "headings": 0,
            "body": 0,
            "figures": 0,
            "empty": 0
        }
        
        # 加载配置
        self.config = None
        if config_path and os.path.exists(config_path):
            self.config = load_config(config_path)
        
        # 默认规范
        self.spec = {
            "cover_title": {"font": "黑体", "size_pt": 22, "bold": True},
            "cover_company": {"font": "宋体", "size_pt": 16, "bold": True},
            "heading1": {"font": "黑体", "size_pt": 16, "bold": False, "first_line_indent_cm": 1.18, "line_spacing": 1.5},
            "heading2": {"font": "黑体", "size_pt": 16, "bold": False, "first_line_indent_cm": 1.18, "line_spacing": 1.5},
            "body": {"font_en": "Times New Roman", "font_cn": "宋体", "size_pt": 16, "first_line_indent_cm": 1.18, "line_spacing": 1.5},
            "figure_caption": {"font": "宋体", "size_pt": 14, "align": "center"}
        }
        
        # 从配置文件更新规范
        if self.config:
            if 'cover' in self.config:
                self.spec['cover_title'] = self.config['cover']['title']
                self.spec['cover_company'] = self.config['cover']['company']
            if 'heading1' in self.config:
                self.spec['heading1'] = self.config['heading1']
            if 'heading2' in self.config:
                self.spec['heading2'] = self.config['heading2']
            if 'body' in self.config:
                self.spec['body'] = self.config['body']
            if 'figure_caption' in self.config:
                self.spec['figure_caption'] = self.config['figure_caption']
    
    def validate(self):
        """执行验证"""
        print(f"=== 验证文档: {self.docx_path} ===\n")
        
        # 1. 检查页面设置
        self._validate_page_setup()
        
        # 2. 检查段落格式
        self._validate_paragraphs()
        
        # 3. 检查样式定义
        self._validate_styles()
        
        # 4. 生成统计
        self._generate_stats()
        
        return len(self.errors) == 0
    
    def _validate_page_setup(self):
        """验证页面设置"""
        if not self.config:
            return
        
        section = self.doc.sections[0]
        page_config = self.config.get('page', {})
        margins_config = page_config.get('margins', {})
        
        # 检查纸张大小
        width_cm = section.page_width / 360000
        height_cm = section.page_height / 360000
        expected_width = page_config.get('width_cm', 21.0)
        expected_height = page_config.get('height_cm', 29.7)
        
        if abs(width_cm - expected_width) > 0.1 or abs(height_cm - expected_height) > 0.1:
            self.errors.append(f"纸张大小不正确: {width_cm:.1f}cm x {height_cm:.1f}cm (应为{expected_width}cm x {expected_height}cm)")
        else:
            print(f"✅ 纸张大小: {width_cm:.1f}cm x {height_cm:.1f}cm")
        
        # 检查页边距
        margins = {
            "上": section.top_margin / 360000,
            "下": section.bottom_margin / 360000,
            "左": section.left_margin / 360000,
            "右": section.right_margin / 360000
        }
        
        expected_margins = {
            "上": margins_config.get('top_cm', 2.54),
            "下": margins_config.get('bottom_cm', 2.54),
            "左": margins_config.get('left_cm', 3.17),
            "右": margins_config.get('right_cm', 3.17)
        }
        
        for name, value in margins.items():
            expected = expected_margins[name]
            if abs(value - expected) > 0.1:
                self.warnings.append(f"{name}边距: {value:.2f}cm (建议{expected}cm)")
            else:
                print(f"✅ {name}边距: {value:.2f}cm")
    
    def _validate_paragraphs(self):
        """验证段落格式"""
        in_cover = True
        
        for i, para in enumerate(self.doc.paragraphs):
            self.stats["paragraphs"] += 1
            
            text = para.text.strip()
            style_name = para.style.name
            
            if not text:
                self.stats["empty"] += 1
                continue
            
            if in_cover and style_name.startswith('Heading'):
                in_cover = False
            
            if in_cover:
                self._validate_cover_paragraph(i, para, text, style_name)
            
            if not in_cover:
                self._validate_body_paragraph(i, para, text, style_name)
    
    def _validate_cover_paragraph(self, index, para, text, style_name):
        """验证封面段落"""
        if style_name == 'Normal' and len(text) > 3:
            for run in para.runs:
                if run.font.size:
                    size_pt = run.font.size.pt
                    expected = self.spec['cover_title']['size_pt']
                    if abs(size_pt - expected) > 1:
                        self.errors.append(f"封面标题字号不正确: {size_pt}pt (应为{expected}pt) [{index}]")
                    else:
                        print(f"✅ 封面标题字号: {size_pt}pt")
                    
                    if self.spec['cover_title']['bold'] and not run.font.bold:
                        self.warnings.append(f"封面标题未加粗 [{index}]")
                    break
        
        if style_name == '公司名称':
            for run in para.runs:
                if run.font.size:
                    size_pt = run.font.size.pt
                    expected = self.spec['cover_company']['size_pt']
                    if abs(size_pt - expected) > 1:
                        self.errors.append(f"公司名称字号不正确: {size_pt}pt (应为{expected}pt) [{index}]")
                    else:
                        print(f"✅ 公司名称字号: {size_pt}pt")
                    break
    
    def _validate_body_paragraph(self, index, para, text, style_name):
        """验证正文段落"""
        if style_name == 'Heading 1':
            self.stats["headings"] += 1
            self._check_heading_format(para, index, 1)
        
        elif style_name == 'Heading 2':
            self.stats["headings"] += 1
            self._check_heading_format(para, index, 2)
        
        elif style_name in ['正文内容', 'Normal']:
            if text.startswith('图') and ('示意' in text or '页面' in text or '界面' in text or '展示' in text):
                self.stats["figures"] += 1
                self._check_figure_caption(para, index)
            else:
                self.stats["body"] += 1
                self._check_body_format(para, index)
    
    def _check_heading_format(self, para, index, level):
        """检查标题格式"""
        spec_key = f"heading{level}"
        spec = self.spec.get(spec_key, {})
        
        for run in para.runs:
            if run.font.size:
                size_pt = run.font.size.pt
                expected = spec.get("size_pt", 16)
                if abs(size_pt - expected) > 1:
                    self.errors.append(f"Heading {level}字号不正确: {size_pt}pt (应为{expected}pt) [{index}]")
            
            if spec.get("bold") is not None:
                if run.font.bold != spec["bold"]:
                    self.warnings.append(f"Heading {level}加粗状态不匹配 [{index}]")
            break
        
        if para.paragraph_format.first_line_indent:
            indent_cm = para.paragraph_format.first_line_indent.cm
            expected = spec.get("first_line_indent_cm", 1.18)
            if abs(indent_cm - expected) > 0.1:
                self.warnings.append(f"Heading {level}首行缩进: {indent_cm:.2f}cm (建议{expected}cm) [{index}]")
        
        if para.paragraph_format.line_spacing:
            line_spacing = para.paragraph_format.line_spacing
            expected = spec.get("line_spacing", 1.5)
            if abs(line_spacing - expected) > 0.1:
                self.warnings.append(f"Heading {level}行距: {line_spacing:.2f} (建议{expected}) [{index}]")
    
    def _check_body_format(self, para, index):
        """检查正文格式"""
        spec = self.spec.get("body", {})
        
        for run in para.runs:
            if run.font.size:
                size_pt = run.font.size.pt
                expected = spec.get("size_pt", 16)
                if abs(size_pt - expected) > 1:
                    self.errors.append(f"正文字号不正确: {size_pt}pt (应为{expected}pt) [{index}]")
            break
        
        if para.paragraph_format.first_line_indent:
            indent_cm = para.paragraph_format.first_line_indent.cm
            expected = spec.get("first_line_indent_cm", 1.18)
            if abs(indent_cm - expected) > 0.1:
                self.warnings.append(f"正文首行缩进: {indent_cm:.2f}cm (建议{expected}cm) [{index}]")
        
        if para.paragraph_format.line_spacing:
            line_spacing = para.paragraph_format.line_spacing
            expected = spec.get("line_spacing", 1.5)
            if abs(line_spacing - expected) > 0.1:
                self.warnings.append(f"正文章距: {line_spacing:.2f} (建议{expected}) [{index}]")
    
    def _check_figure_caption(self, para, index):
        """检查图名格式"""
        spec = self.spec.get("figure_caption", {})
        
        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            self.warnings.append(f"图名未居中 [{index}]")
        
        for run in para.runs:
            if run.font.size:
                size_pt = run.font.size.pt
                expected = spec.get("size_pt", 14)
                if abs(size_pt - expected) > 1:
                    self.errors.append(f"图名字号不正确: {size_pt}pt (应为{expected}pt) [{index}]")
            break
    
    def _validate_styles(self):
        """验证样式定义"""
        required_styles = ['Heading 1', 'Heading 2', '正文内容']
        
        for style_name in required_styles:
            try:
                style = self.doc.styles[style_name]
                print(f"✅ 样式 '{style_name}' 存在")
            except KeyError:
                self.warnings.append(f"缺少样式定义: {style_name}")
    
    def _generate_stats(self):
        """生成统计信息"""
        print(f"\n=== 文档统计 ===")
        print(f"总段落数: {self.stats['paragraphs']}")
        print(f"标题数: {self.stats['headings']}")
        print(f"正文段落数: {self.stats['body']}")
        print(f"图片说明数: {self.stats['figures']}")
        print(f"空段落数: {self.stats['empty']}")
        
        total_chars = sum(len(p.text) for p in self.doc.paragraphs)
        print(f"总字数: {total_chars}")
    
    def get_report(self):
        """获取验证报告"""
        return {
            "file": self.docx_path,
            "valid": len(self.errors) == 0,
            "errors": self.errors,
            "warnings": self.warnings,
            "stats": self.stats,
            "total_chars": sum(len(p.text) for p in self.doc.paragraphs)
        }
    
    def fix_errors(self):
        """自动修复错误"""
        print("\n=== 自动修复 ===\n")
        
        if not self.config:
            print("⚠️ 未加载配置文件，无法自动修复")
            return 0
        
        fixed = 0
        body_config = self.config.get('body', {})
        h1_config = self.config.get('heading1', {})
        h2_config = self.config.get('heading2', {})
        figure_config = self.config.get('figure_caption', {})
        
        def apply_format(para, style_config):
            """应用格式"""
            nonlocal fixed
            
            pf = para.paragraph_format
            
            if 'first_line_indent_cm' in style_config and style_config['first_line_indent_cm'] is not None:
                if pf.first_line_indent is None or abs(pf.first_line_indent.cm - style_config['first_line_indent_cm']) > 0.1:
                    pf.first_line_indent = Cm(style_config['first_line_indent_cm'])
                    fixed += 1
            
            if 'line_spacing' in style_config and style_config['line_spacing'] is not None:
                if pf.line_spacing is None or abs(pf.line_spacing - style_config['line_spacing']) > 0.1:
                    pf.line_spacing = style_config['line_spacing']
                    fixed += 1
            
            if 'space_before_pt' in style_config and style_config['space_before_pt'] is not None:
                pf.space_before = Pt(style_config['space_before_pt'])
            
            if 'space_after_pt' in style_config and style_config['space_after_pt'] is not None:
                pf.space_after = Pt(style_config['space_after_pt'])
            
            if 'align' in style_config:
                align_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if style_config['align'] in align_map:
                    pf.alignment = align_map[style_config['align']]
            
            for run in para.runs:
                font_en = style_config.get('font_en', style_config.get('font'))
                font_cn = style_config.get('font_cn', style_config.get('font'))
                size_pt = style_config.get('size_pt')
                bold = style_config.get('bold')
                
                if font_en:
                    run.font.name = font_en
                if size_pt:
                    run.font.size = Pt(size_pt)
                if bold is not None:
                    run.font.bold = bold
                
                if font_cn:
                    r = run._element
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), font_cn)
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name
            
            if style_name == '正文内容':
                apply_format(para, body_config)
            elif style_name == 'Heading 1':
                apply_format(para, h1_config)
            elif style_name == 'Heading 2':
                apply_format(para, h2_config)
            elif text.startswith('图') and ('示意' in text or '页面' in text or '界面' in text):
                apply_format(para, figure_config)
        
        print(f"共修复 {fixed} 个问题")
        return fixed
    
    def save(self, output_path):
        """保存修复后的文档"""
        self.doc.save(output_path)
        print(f"\n✅ 已保存: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="软著文档格式验证")
    parser.add_argument("--docx", required=True, help="要验证的Word文档")
    parser.add_argument("--config", default="config/format_spec.json", help="格式规范配置文件")
    parser.add_argument("--fix", action="store_true", help="自动修复错误")
    parser.add_argument("--report", help="输出JSON报告路径")
    args = parser.parse_args()
    
    if not os.path.exists(args.docx):
        print(f"❌ 文件不存在: {args.docx}")
        exit(1)
    
    # 创建验证器
    validator = FormatValidator(args.docx, args.config)
    
    # 执行验证
    valid = validator.validate()
    
    # 输出结果
    print(f"\n=== 验证结果 ===")
    if valid:
        print("✅ 文档格式验证通过")
    else:
        print(f"❌ 发现 {len(validator.errors)} 个错误")
        for error in validator.errors:
            print(f"   - {error}")
    
    if validator.warnings:
        print(f"\n⚠️ {len(validator.warnings)} 个警告")
        for warning in validator.warnings[:5]:
            print(f"   - {warning}")
        if len(validator.warnings) > 5:
            print(f"   ... 还有 {len(validator.warnings) - 5} 个警告")
    
    # 自动修复
    if args.fix and not valid:
        fixed = validator.fix_errors()
        if fixed > 0:
            output_path = args.docx.replace(".docx", "_fixed.docx")
            validator.save(output_path)
            
            print("\n=== 重新验证 ===")
            validator2 = FormatValidator(output_path, args.config)
            valid2 = validator2.validate()
            
            if valid2:
                print("✅ 修复后文档格式验证通过")
            else:
                print(f"❌ 仍有 {len(validator2.errors)} 个错误")
    
    # 输出JSON报告
    if args.report:
        report = validator.get_report()
        with open(args.report, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        print(f"\n📄 报告已保存: {args.report}")
    
    exit(0 if valid else 1)


if __name__ == "__main__":
    main()
