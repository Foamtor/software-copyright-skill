#!/usr/bin/env python3
"""
Word文档后处理脚本

功能：
1. 分割长段落（每个换行符分割成独立段落）
2. 分割图名（将图名分割成独立段落）
3. 插入图片（在图名上方插入实际图片）
4. 删除连续空行

用法：
  python3 post_process.py \
    --input input.docx \
    --output output.docx \
    --content content.json \
    --image-dir images/
"""

import argparse
import json
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少依赖：pip install python-docx")
    exit(1)


def split_long_paragraphs(doc):
    """将长段落分割成多个独立段落"""
    paragraphs = list(doc.paragraphs)
    split_count = 0
    
    for i in range(len(paragraphs)-1, -1, -1):
        para = paragraphs[i]
        text = para.text
        
        if '\n' in text:
            lines = text.split('\n')
            main_text = lines[0]
            other_lines = [line for line in lines[1:] if line.strip()]
            
            if other_lines:
                # 清空原段落
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = main_text
                else:
                    para.add_run(main_text)
                
                # 在原段落后插入新段落
                prev_para = para
                for line in other_lines:
                    new_p = OxmlElement('w:p')
                    prev_para._element.addnext(new_p)
                    from docx.text.paragraph import Paragraph
                    new_para = Paragraph(new_p, para._parent)
                    
                    # 复制原段落的格式
                    new_para.style = para.style
                    new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                    new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                    new_para.paragraph_format.alignment = para.paragraph_format.alignment
                    
                    # 设置文本
                    run = new_para.add_run(line)
                    
                    # 复制字体格式
                    if para.runs:
                        ref_run = para.runs[0]
                        run.font.name = ref_run.font.name
                        run.font.size = ref_run.font.size
                        run.font.bold = ref_run.font.bold
                        # 设置中文字体
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.append(rFonts)
                        rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    prev_para = new_para
                    split_count += 1
    
    return split_count


def split_paragraph_at_figure_captions(doc):
    """将图名分割成独立段落"""
    paragraphs = list(doc.paragraphs)
    split_count = 0
    
    for i in range(len(paragraphs)-1, -1, -1):
        para = paragraphs[i]
        text = para.text
        
        # 查找图名
        fig_pattern = r'(图\d+\s+[^\n]+)'
        matches = list(re.finditer(fig_pattern, text))
        
        if len(matches) > 0:
            last_end = len(text)
            fig_captions = []
            
            for match in reversed(matches):
                fig_captions.insert(0, match.group(1))
                last_end = match.start()
            
            main_text = text[:last_end].strip()
            
            if main_text:
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = main_text
                else:
                    para.add_run(main_text)
                
                prev_para = para
                for fig_caption in fig_captions:
                    new_p = OxmlElement('w:p')
                    prev_para._element.addnext(new_p)
                    from docx.text.paragraph import Paragraph
                    new_para = Paragraph(new_p, para._parent)
                    
                    # 设置图名格式：14pt，居中
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = new_para.add_run(fig_caption)
                    run.font.size = Pt(14)
                    
                    prev_para = new_para
                    split_count += 1
    
    return split_count


def build_image_mapping(content, image_dir):
    """构建图片映射：图号 -> 图片路径"""
    image_mapping = {}
    
    for section in content.get('sections', []):
        for img in section.get('images', []):
            ref = img.get('ref', '')
            path = img.get('path', '')
            if ref and path:
                full_path = os.path.join(image_dir, path)
                if os.path.exists(full_path):
                    image_mapping[ref] = full_path
        
        for sub in section.get('subsections', []):
            for img in sub.get('images', []):
                ref = img.get('ref', '')
                path = img.get('path', '')
                if ref and path:
                    full_path = os.path.join(image_dir, path)
                    if os.path.exists(full_path):
                        image_mapping[ref] = full_path
    
    return image_mapping


def insert_images_before_caption(doc, image_mapping):
    """在图名上方插入图片"""
    paragraphs = list(doc.paragraphs)
    insert_count = 0
    inserted_refs = set()
    
    for i in range(len(paragraphs)-1, -1, -1):
        para = paragraphs[i]
        text = para.text.strip()
        
        # 检查是否是图名段落
        match = re.match(r'^(图\d+)\s+', text)
        if match:
            fig_ref = match.group(1)
            if fig_ref in image_mapping and fig_ref not in inserted_refs:
                img_path = image_mapping[fig_ref]
                
                # 在图名段落前插入图片
                new_p = OxmlElement('w:p')
                para._element.addprevious(new_p)
                from docx.text.paragraph import Paragraph
                new_para = Paragraph(new_p, para._parent)
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = new_para.add_run()
                run.add_picture(img_path, width=Inches(5.0))
                
                insert_count += 1
                inserted_refs.add(fig_ref)
    
    return insert_count


def remove_consecutive_empty_lines(doc):
    """删除连续空行，只保留一个"""
    paragraphs = list(doc.paragraphs)
    to_remove = []
    
    for i in range(len(paragraphs) - 1):
        curr_empty = not paragraphs[i].text.strip()
        next_empty = not paragraphs[i+1].text.strip()
        
        if curr_empty and next_empty:
            to_remove.append(i+1)
    
    # 从后往前删除
    for i in reversed(to_remove):
        p = paragraphs[i]._element
        p.getparent().remove(p)
    
    return len(to_remove)


def post_process(input_path, output_path, content_path=None, image_dir=None):
    """后处理文档"""
    print(f"📄 读取文档: {input_path}")
    doc = Document(input_path)
    
    # 1. 分割长段落
    print("\n📝 分割长段落...")
    split_count = split_long_paragraphs(doc)
    print(f"   分割了 {split_count} 个段落")
    
    # 2. 分割图名
    print("\n📝 分割图名...")
    fig_split_count = split_paragraph_at_figure_captions(doc)
    print(f"   分割了 {fig_split_count} 个图名")
    
    # 3. 插入图片
    if content_path and image_dir:
        print(f"\n📷 插入图片...")
        with open(content_path, 'r', encoding='utf-8') as f:
            content = json.load(f)
        
        image_mapping = build_image_mapping(content, image_dir)
        print(f"   图片映射: {len(image_mapping)}张")
        
        insert_count = insert_images_before_caption(doc, image_mapping)
        print(f"   插入了 {insert_count} 张图片")
    
    # 4. 删除连续空行
    print("\n🧹 删除连续空行...")
    removed_count = remove_consecutive_empty_lines(doc)
    print(f"   删除了 {removed_count} 个空行")
    
    # 保存文档
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    
    print(f"\n✅ 后处理完成: {output_path}")
    
    return {
        'split_paragraphs': split_count,
        'split_figures': fig_split_count,
        'removed_empty_lines': removed_count
    }


def main():
    parser = argparse.ArgumentParser(description="Word文档后处理")
    parser.add_argument("--input", required=True, help="输入文档路径")
    parser.add_argument("--output", required=True, help="输出文档路径")
    parser.add_argument("--content", help="content.json路径（用于插入图片）")
    parser.add_argument("--image-dir", help="图片目录路径")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"❌ 输入文件不存在: {args.input}")
        exit(1)
    
    post_process(args.input, args.output, args.content, args.image_dir)


if __name__ == "__main__":
    main()
