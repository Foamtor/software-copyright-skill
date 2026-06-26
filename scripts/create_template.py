#!/usr/bin/env python3
"""
从参考文档自动创建Word模板。

用法：
  python3 create_template.py --ref 参考文档.docx --output templates/manual_template.docx

功能：
  1. 分析参考文档结构（封面、章节、正文）
  2. 自动替换内容为Jinja2占位符
  3. 保留原有格式（字体、字号、行距、缩进等）
  4. 生成可用的docxtpl模板
"""

import argparse
import json
import os
import re
import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 缺少依赖：pip install python-docx")
    exit(1)


class TemplateCreator:
    """模板创建器"""
    
    def __init__(self, ref_path):
        """初始化"""
        self.ref_path = ref_path
        self.doc = Document(ref_path)
        self.analysis = {
            "cover": {},
            "sections": [],
            "styles": {}
        }
    
    def analyze(self):
        """分析文档结构"""
        print("=== 分析文档结构 ===\n")
        
        # 1. 提取页面设置
        self._extract_page_setup()
        
        # 2. 提取样式定义
        self._extract_styles()
        
        # 3. 分析段落结构
        self._analyze_paragraphs()
        
        return self.analysis
    
    def _extract_page_setup(self):
        """提取页面设置"""
        section = self.doc.sections[0]
        self.analysis["page"] = {
            "width_cm": round(section.page_width / 360000, 1),
            "height_cm": round(section.page_height / 360000, 1),
            "margin_top_cm": round(section.top_margin / 360000, 2),
            "margin_bottom_cm": round(section.bottom_margin / 360000, 2),
            "margin_left_cm": round(section.left_margin / 360000, 2),
            "margin_right_cm": round(section.right_margin / 360000, 2),
            "header_distance_cm": round(section.header_distance / 360000, 2),
            "footer_distance_cm": round(section.footer_distance / 360000, 2)
        }
        print(f"✅ 页面设置：A4 ({self.analysis['page']['width_cm']}cm x {self.analysis['page']['height_cm']}cm)")
    
    def _extract_styles(self):
        """提取样式定义"""
        style_names = ['Heading 1', 'Heading 2', 'Heading 3', '正文内容', '公司名称', 'Normal']
        
        for name in style_names:
            try:
                style = self.doc.styles[name]
                style_info = {
                    "font": style.font.name,
                    "size_pt": round(style.font.size.pt, 1) if style.font.size else None,
                    "bold": style.font.bold,
                    "line_spacing": style.paragraph_format.line_spacing,
                    "space_before_pt": round(style.paragraph_format.space_before.pt, 1) if style.paragraph_format.space_before else None,
                    "space_after_pt": round(style.paragraph_format.space_after.pt, 1) if style.paragraph_format.space_after else None,
                    "first_line_indent_cm": round(style.paragraph_format.first_line_indent.cm, 2) if style.paragraph_format.first_line_indent else None
                }
                self.analysis["styles"][name] = style_info
                print(f"✅ 样式 '{name}': {style_info['font'] or '继承'}, {style_info['size_pt']}pt")
            except KeyError:
                pass
    
    def _analyze_paragraphs(self):
        """分析段落结构"""
        in_cover = True
        cover_title_idx = None
        cover_subtitle_idx = None
        cover_company_idx = None
        cover_date_idx = None
        current_section = None
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name
            
            # 跳过空段落
            if not text:
                continue
            
            # 识别封面元素
            if in_cover:
                # 软件标题（通常是最大的居中文字）
                if style_name == 'Normal' and len(text) > 3:
                    if cover_title_idx is None:
                        cover_title_idx = i
                        self.analysis["cover"]["title"] = {
                            "index": i,
                            "text": text,
                            "style": style_name
                        }
                        print(f"✅ 封面标题 [{i}]: {text[:30]}")
                    elif cover_subtitle_idx is None and len(text) < 20:
                        cover_subtitle_idx = i
                        self.analysis["cover"]["subtitle"] = {
                            "index": i,
                            "text": text,
                            "style": style_name
                        }
                        print(f"✅ 封面副标题 [{i}]: {text[:30]}")
                
                # 公司名称
                if style_name == '公司名称':
                    if '公司' in text or '中心' in text or '局' in text or '院' in text:
                        cover_company_idx = i
                        self.analysis["cover"]["company"] = {
                            "index": i,
                            "text": text,
                            "style": style_name
                        }
                        print(f"✅ 公司名称 [{i}]: {text[:30]}")
                    elif '年' in text and '月' in text:
                        cover_date_idx = i
                        self.analysis["cover"]["date"] = {
                            "index": i,
                            "text": text,
                            "style": style_name
                        }
                        print(f"✅ 日期 [{i}]: {text[:30]}")
                
                # 检测是否进入正文（遇到Heading样式）
                if style_name.startswith('Heading'):
                    in_cover = False
            
            # 识别章节
            if not in_cover:
                if style_name == 'Heading 1':
                    current_section = {
                        "heading_index": i,
                        "heading_text": text,
                        "content_indices": [],
                        "subsections": []
                    }
                    self.analysis["sections"].append(current_section)
                elif style_name == 'Heading 2' and current_section:
                    current_section["subsections"].append({
                        "heading_index": i,
                        "heading_text": text,
                        "content_indices": []
                    })
                elif style_name in ['正文内容', 'Normal'] and current_section:
                    if current_section["subsections"]:
                        current_section["subsections"][-1]["content_indices"].append(i)
                    else:
                        current_section["content_indices"].append(i)
        
        print(f"\n✅ 识别到 {len(self.analysis['sections'])} 个一级章节")
        for sec in self.analysis["sections"][:3]:
            print(f"   - {sec['heading_text'][:30]}")
    
    def create_template(self, output_path):
        """创建模板"""
        print("\n=== 创建模板 ===\n")
        
        # 复制参考文档
        shutil.copy2(self.ref_path, output_path)
        doc = Document(output_path)
        
        # 替换封面内容
        self._replace_cover(doc)
        
        # 替换章节内容
        self._replace_sections(doc)
        
        # 保存
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        
        print(f"\n✅ 模板已创建: {output_path}")
        return output_path
    
    def _replace_cover(self, doc):
        """替换封面内容为Jinja2占位符"""
        cover = self.analysis.get("cover", {})
        
        # 替换标题
        if "title" in cover:
            idx = cover["title"]["index"]
            para = doc.paragraphs[idx]
            self._clear_and_set_placeholder(para, "{{cover_title}}")
            print(f"✅ 替换封面标题 [{idx}]")
        
        # 替换副标题
        if "subtitle" in cover:
            idx = cover["subtitle"]["index"]
            para = doc.paragraphs[idx]
            self._clear_and_set_placeholder(para, "{{cover_subtitle}}")
            print(f"✅ 替换封面副标题 [{idx}]")
        
        # 替换公司名称
        if "company" in cover:
            idx = cover["company"]["index"]
            para = doc.paragraphs[idx]
            self._clear_and_set_placeholder(para, "{{cover_company}}")
            print(f"✅ 替换公司名称 [{idx}]")
        
        # 替换日期
        if "date" in cover:
            idx = cover["date"]["index"]
            para = doc.paragraphs[idx]
            self._clear_and_set_placeholder(para, "{{cover_date}}")
            print(f"✅ 替换日期 [{idx}]")
    
    def _replace_sections(self, doc):
        """替换章节内容为Jinja2占位符"""
        sections = self.analysis.get("sections", [])
        
        if not sections:
            print("⚠️ 未识别到章节，跳过章节替换")
            return
        
        # 记录需要替换的段落索引
        replace_map = {}
        
        for sec_idx, section in enumerate(sections):
            # 替换一级标题
            h_idx = section["heading_index"]
            replace_map[h_idx] = {
                "type": "heading1",
                "placeholder": "{{section.heading}}"
            }
            
            # 替换正文内容
            for c_idx in section["content_indices"]:
                replace_map[c_idx] = {
                    "type": "content",
                    "placeholder": "{{section.content}}"
                }
            
            # 替换子章节
            for sub_idx, sub in enumerate(section["subsections"]):
                # 子标题
                replace_map[sub["heading_index"]] = {
                    "type": "heading2",
                    "placeholder": "{{sub.heading}}"
                }
                
                # 子正文
                for c_idx in sub["content_indices"]:
                    replace_map[c_idx] = {
                        "type": "sub_content",
                        "placeholder": "{{sub.content}}"
                    }
        
        # 执行替换（只替换第一个章节作为模板示例）
        if sections:
            first_section = sections[0]
            
            # 替换第一个一级标题
            h_idx = first_section["heading_index"]
            para = doc.paragraphs[h_idx]
            self._clear_and_set_placeholder(para, "{{section.heading}}")
            print(f"✅ 替换一级标题 [{h_idx}]")
            
            # 替换第一个正文段落
            if first_section["content_indices"]:
                c_idx = first_section["content_indices"][0]
                para = doc.paragraphs[c_idx]
                self._clear_and_set_placeholder(para, "{{section.content}}")
                print(f"✅ 替换正文 [{c_idx}]")
            
            # 替换第一个子章节
            if first_section["subsections"]:
                sub = first_section["subsections"][0]
                
                # 子标题
                para = doc.paragraphs[sub["heading_index"]]
                self._clear_and_set_placeholder(para, "{{sub.heading}}")
                print(f"✅ 替换子标题 [{sub['heading_index']}]")
                
                # 子正文
                if sub["content_indices"]:
                    para = doc.paragraphs[sub["content_indices"][0]]
                    self._clear_and_set_placeholder(para, "{{sub.content}}")
                    print(f"✅ 替换子正文 [{sub['content_indices'][0]}]")
    
    def _clear_and_set_placeholder(self, para, placeholder):
        """清空段落并设置占位符（保留格式）"""
        # 保存第一个run的格式
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold
            font_italic = first_run.font.italic
            font_color = first_run.font.color.rgb if first_run.font.color else None
        else:
            font_name = None
            font_size = None
            font_bold = None
            font_italic = None
            font_color = None
        
        # 清空所有run
        for run in para.runs:
            run.text = ""
        
        # 设置占位符
        if para.runs:
            para.runs[0].text = placeholder
            # 恢复格式
            if font_name:
                para.runs[0].font.name = font_name
            if font_size:
                para.runs[0].font.size = font_size
            if font_bold is not None:
                para.runs[0].font.bold = font_bold
            if font_italic is not None:
                para.runs[0].font.italic = font_italic
            if font_color:
                para.runs[0].font.color.rgb = font_color
        else:
            run = para.add_run(placeholder)
    
    def get_analysis_summary(self):
        """获取分析摘要"""
        return {
            "page": self.analysis.get("page", {}),
            "styles": self.analysis.get("styles", {}),
            "cover": {
                "has_title": "title" in self.analysis.get("cover", {}),
                "has_company": "company" in self.analysis.get("cover", {}),
                "has_date": "date" in self.analysis.get("cover", {})
            },
            "sections_count": len(self.analysis.get("sections", [])),
            "total_paragraphs": len(self.doc.paragraphs)
        }


def main():
    parser = argparse.ArgumentParser(description="从参考文档自动创建Word模板")
    parser.add_argument("--ref", required=True, help="参考文档路径")
    parser.add_argument("--output", required=True, help="输出模板路径")
    parser.add_argument("--analyze-only", action="store_true", help="仅分析，不创建模板")
    parser.add_argument("--json", action="store_true", help="输出JSON格式分析结果")
    args = parser.parse_args()
    
    if not os.path.exists(args.ref):
        print(f"❌ 参考文档不存在: {args.ref}")
        exit(1)
    
    # 创建模板创建器
    creator = TemplateCreator(args.ref)
    
    # 分析文档
    analysis = creator.analyze()
    
    if args.json:
        # 输出JSON格式
        summary = creator.get_analysis_summary()
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    
    if args.analyze_only:
        print("\n✅ 分析完成（仅分析模式）")
        return
    
    # 创建模板
    creator.create_template(args.output)
    
    print("\n=== 使用说明 ===")
    print("1. 用Word打开生成的模板")
    print("2. 检查Jinja2占位符是否正确")
    print("3. 根据需要添加循环和条件标签")
    print("4. 保存模板")
    print("\n示例Jinja2标签：")
    print("  {{cover_title}}        - 封面标题")
    print("  {{section.heading}}    - 章节标题")
    print("  {{section.content}}    - 章节内容")
    print("  {% for img in section.images %}")
    print("  {{img.ref}} {{img.description}}")
    print("  {% endfor %}")


if __name__ == "__main__":
    main()
