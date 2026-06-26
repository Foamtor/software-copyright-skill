#!/usr/bin/env python3
"""
创建符合软著规范的Word模板（使用配置文件）。

用法：
  python3 create_standard_template.py [--config config/format_spec.json] [--output templates/manual_template.docx]

功能：
  1. 读取格式规范配置文件
  2. 按规范创建Word模板
  3. 设置所有样式（标题、正文、图名等）
  4. 添加Jinja2占位符
"""

import argparse
import json
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def load_config(config_path):
    """加载格式规范配置"""
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def set_font(run, font_name, font_size_pt, bold=False, color="000000"):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_format(para, config):
    """设置段落格式"""
    pf = para.paragraph_format
    
    if 'first_line_indent_cm' in config and config['first_line_indent_cm'] is not None:
        pf.first_line_indent = Cm(config['first_line_indent_cm'])
    
    if 'line_spacing' in config and config['line_spacing'] is not None:
        pf.line_spacing = config['line_spacing']
    
    if 'space_before_pt' in config and config['space_before_pt'] is not None:
        pf.space_before = Pt(config['space_before_pt'])
    
    if 'space_after_pt' in config and config['space_after_pt'] is not None:
        pf.space_after = Pt(config['space_after_pt'])
    
    if 'align' in config:
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if config['align'] in align_map:
            pf.alignment = align_map[config['align']]


def setup_style(doc, style_name, config):
    """设置样式"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    
    # 设置字体
    if 'font' in config:
        style.font.name = config['font']
    if 'font_en' in config:
        style.font.name = config['font_en']
    if 'size_pt' in config:
        style.font.size = Pt(config['size_pt'])
    if 'bold' in config:
        style.font.bold = config['bold']
    if 'color' in config:
        style.font.color.rgb = RGBColor.from_string(config['color'])
    
    # 设置段落格式
    pf = style.paragraph_format
    if 'first_line_indent_cm' in config and config['first_line_indent_cm'] is not None:
        pf.first_line_indent = Cm(config['first_line_indent_cm'])
    if 'line_spacing' in config and config['line_spacing'] is not None:
        pf.line_spacing = config['line_spacing']
    if 'space_before_pt' in config and config['space_before_pt'] is not None:
        pf.space_before = Pt(config['space_before_pt'])
    if 'space_after_pt' in config and config['space_after_pt'] is not None:
        pf.space_after = Pt(config['space_after_pt'])
    if 'align' in config:
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if config['align'] in align_map:
            pf.alignment = align_map[config['align']]
    
    # 设置中文字体
    font_cn = config.get('font_cn', config.get('font'))
    if font_cn:
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_cn)
    
    return style


def create_template(config, output_path):
    """创建模板"""
    doc = Document()
    
    # 1. 设置页面
    page_config = config['page']
    section = doc.sections[0]
    section.page_width = Cm(page_config['width_cm'])
    section.page_height = Cm(page_config['height_cm'])
    section.top_margin = Cm(page_config['margins']['top_cm'])
    section.bottom_margin = Cm(page_config['margins']['bottom_cm'])
    section.left_margin = Cm(page_config['margins']['left_cm'])
    section.right_margin = Cm(page_config['margins']['right_cm'])
    section.header_distance = Cm(page_config['margins']['header_distance_cm'])
    section.footer_distance = Cm(page_config['margins']['footer_distance_cm'])
    
    # 2. 设置页眉
    header_config = config['header']
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(header_config['left']['text'])
    set_font(run, header_config['left']['font'], header_config['left']['size_pt'])
    
    # 3. 设置样式
    setup_style(doc, 'Heading 1', config['heading1'])
    setup_style(doc, 'Heading 2', config['heading2'])
    setup_style(doc, 'Heading 3', config.get('heading3', config['heading2']))
    setup_style(doc, '正文内容', config['body'])
    setup_style(doc, '公司名称', config['company_name'])
    
    # 4. 添加封面内容
    cover_config = config['cover']
    
    # 软件名称
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("{{cover_title}}")
    set_font(run, cover_config['title']['font'], cover_config['title']['size_pt'], 
             bold=cover_config['title']['bold'])
    
    # 副标题
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("{{cover_subtitle}}")
    set_font(run, cover_config['subtitle']['font'], cover_config['subtitle']['size_pt'],
             bold=cover_config['subtitle'].get('bold', False))
    
    # 空行
    for _ in range(8):
        doc.add_paragraph()
    
    # 公司名称
    company_para = doc.add_paragraph(style='公司名称')
    run = company_para.add_run("{{cover_company}}")
    set_font(run, cover_config['company']['font'], cover_config['company']['size_pt'],
             bold=cover_config['company']['bold'])
    
    # 日期
    date_para = doc.add_paragraph(style='公司名称')
    run = date_para.add_run("{{cover_date}}")
    set_font(run, cover_config['date']['font'], cover_config['date']['size_pt'],
             bold=cover_config['date']['bold'])
    
    # 5. 分页
    doc.add_page_break()
    
    # 6. 添加章节占位
    # 一级章节开始
    h1_start = doc.add_paragraph(style='Heading 1')
    run = h1_start.add_run("{% for section in sections %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 一级标题
    h1_para = doc.add_paragraph(style='Heading 1')
    run = h1_para.add_run("{{section.heading}}")
    set_font(run, config['heading1']['font'], config['heading1']['size_pt'])
    
    # 正文内容
    body_para = doc.add_paragraph(style='正文内容')
    run = body_para.add_run("{{section.content}}")
    set_font(run, config['body']['font_en'], config['body']['size_pt'])
    
    # 图片循环开始
    img_start = doc.add_paragraph()
    run = img_start.add_run("{% for img in section.images %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 图片说明
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run("{{img.ref}} {{img.description}}")
    set_font(run, config['figure_caption']['font_en'], config['figure_caption']['size_pt'])
    
    # 图片循环结束
    img_end = doc.add_paragraph()
    run = img_end.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 注意事项循环开始
    note_start = doc.add_paragraph(style='正文内容')
    run = note_start.add_run("{% for note in section.notes %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 注意事项
    note_para = doc.add_paragraph(style='正文内容')
    run = note_para.add_run("注意：{{note}}")
    set_font(run, config['body']['font_en'], config['body']['size_pt'])
    
    # 注意事项循环结束
    note_end = doc.add_paragraph()
    run = note_end.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 子章节循环开始
    sub_start = doc.add_paragraph()
    run = sub_start.add_run("{% for sub in section.subsections %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 子标题
    h2_para = doc.add_paragraph(style='Heading 2')
    run = h2_para.add_run("{{sub.heading}}")
    set_font(run, config['heading2']['font'], config['heading2']['size_pt'])
    
    # 子正文
    sub_body = doc.add_paragraph(style='正文内容')
    run = sub_body.add_run("{{sub.content}}")
    set_font(run, config['body']['font_en'], config['body']['size_pt'])
    
    # 子章节循环结束
    sub_end = doc.add_paragraph()
    run = sub_end.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 章节循环结束
    section_end = doc.add_paragraph()
    run = section_end.add_run("{% endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 保存
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ 模板已创建: {output_path}")
    print(f"   配置文件: {config_path}")
    
    # 验证格式
    print("\n=== 格式验证 ===")
    doc2 = Document(output_path)
    for style_name in ['Heading 1', 'Heading 2', '正文内容', '公司名称']:
        style = doc2.styles[style_name]
        print(f"{style_name}: 字体={style.font.name}, 字号={style.font.size}, 加粗={style.font.bold}")
    
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="创建符合软著规范的Word模板")
    parser.add_argument("--config", default="config/format_spec.json", help="格式规范配置文件路径")
    parser.add_argument("--output", default="templates/manual_template.docx", help="输出模板路径")
    args = parser.parse_args()
    
    # 检查配置文件
    if not os.path.exists(args.config):
        print(f"❌ 配置文件不存在: {args.config}")
        exit(1)
    
    # 加载配置
    config = load_config(args.config)
    config_path = args.config
    
    # 创建模板
    create_template(config, args.output)
