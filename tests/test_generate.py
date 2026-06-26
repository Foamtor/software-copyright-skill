#!/usr/bin/env python3
"""
TDD测试套件 - 软著文档生成Skill

测试层次：
1. 单元测试 - 测试utils.py中的函数
2. 集成测试 - 测试generate_manual.py的完整流程
3. 验收测试 - 验证生成的文档符合Brand Profile规范
"""

import json
import os
import sys
import tempfile
from pathlib import Path

# 添加scripts目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent / 'scripts'))

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 缺少依赖：pip install python-docx")
    exit(1)

from utils import (
    load_config, load_content, validate_content, load_brand_profile,
    count_chars, extract_figure_refs, has_figure_caption,
    get_style_config, apply_brand_style, brand_profile_to_format_spec
)


# ============================================================================
# 测试工具
# ============================================================================

class TestResult:
    """测试结果"""
    def __init__(self, name):
        self.name = name
        self.passed = 0
        self.failed = 0
        self.errors = []
    
    def assert_equal(self, actual, expected, message=""):
        if actual == expected:
            self.passed += 1
        else:
            self.failed += 1
            self.errors.append(f"{message}: 期望 {expected}, 实际 {actual}")
    
    def assert_true(self, condition, message=""):
        if condition:
            self.passed += 1
        else:
            self.failed += 1
            self.errors.append(f"{message}: 条件为False")
    
    def assert_false(self, condition, message=""):
        if not condition:
            self.passed += 1
        else:
            self.failed += 1
            self.errors.append(f"{message}: 条件为True")
    
    def summary(self):
        total = self.passed + self.failed
        status = "✅ PASS" if self.failed == 0 else "❌ FAIL"
        print(f"\n{status} {self.name}: {self.passed}/{total} 通过")
        if self.errors:
            for error in self.errors:
                print(f"   ❌ {error}")


# ============================================================================
# 单元测试 - utils.py
# ============================================================================

def test_count_chars():
    """测试字数统计函数"""
    result = TestResult("count_chars")
    
    # 测试正常文本
    result.assert_equal(count_chars("你好世界"), 4, "中文字符")
    result.assert_equal(count_chars("hello world"), 10, "英文字符（不含空格）")
    result.assert_equal(count_chars("你好 世界"), 4, "中文字符（含空格）")
    result.assert_equal(count_chars(""), 0, "空字符串")
    result.assert_equal(count_chars("  \n  "), 0, "空白字符")
    
    result.summary()
    return result


def test_extract_figure_refs():
    """测试图号提取函数"""
    result = TestResult("extract_figure_refs")
    
    # 测试正常图号
    result.assert_equal(extract_figure_refs("图8 系统总体架构图"), ["图8"], "单个图号")
    result.assert_equal(extract_figure_refs("图8 图9 图10"), ["图8", "图9", "图10"], "多个图号")
    result.assert_equal(extract_figure_refs("没有图号"), [], "没有图号")
    result.assert_equal(extract_figure_refs(""), [], "空字符串")
    
    result.summary()
    return result


def test_has_figure_caption():
    """测试图名检测函数"""
    result = TestResult("has_figure_caption")
    
    # 测试图名检测
    result.assert_true(has_figure_caption("图8 系统总体架构图"), "有图名")
    result.assert_true(has_figure_caption("图1 首页"), "有图名（短）")
    result.assert_false(has_figure_caption("没有图名"), "没有图名")
    result.assert_false(has_figure_caption(""), "空字符串")
    result.assert_false(has_figure_caption("图"), "只有图字")
    
    result.summary()
    return result


def test_validate_content():
    """测试内容验证函数"""
    result = TestResult("validate_content")
    
    # 测试有效内容
    valid_content = {
        "software_name": "测试软件",
        "sections": [
            {
                "heading": "一、测试章节",
                "content": "测试内容"
            }
        ]
    }
    errors = validate_content(valid_content)
    result.assert_equal(len(errors), 0, "有效内容无错误")
    
    # 测试缺少sections
    invalid_content1 = {}
    errors = validate_content(invalid_content1)
    result.assert_true(len(errors) > 0, "缺少sections应报错")
    
    # 测试缺少heading
    invalid_content2 = {
        "sections": [{"content": "测试内容"}]
    }
    errors = validate_content(invalid_content2)
    result.assert_true(len(errors) > 0, "缺少heading应报错")
    
    result.summary()
    return result


def test_brand_profile_loading():
    """测试Brand Profile加载"""
    result = TestResult("brand_profile_loading")
    
    profile_path = Path(__file__).parent.parent / 'config' / 'brand_profile.json'
    if profile_path.exists():
        profile = load_brand_profile(str(profile_path))
        
        # 检查必要字段
        result.assert_true('brand_info' in profile, "有brand_info字段")
        result.assert_true('styles' in profile, "有styles字段")
        result.assert_true('content_rules' in profile, "有content_rules字段")
        
        # 检查样式定义
        styles = profile.get('styles', {})
        result.assert_true('cover_title' in styles, "有cover_title样式")
        result.assert_true('heading1' in styles, "有heading1样式")
        result.assert_true('body' in styles, "有body样式")
        result.assert_true('figure_caption' in styles, "有figure_caption样式")
    else:
        result.assert_true(False, f"Brand Profile文件不存在: {profile_path}")
    
    result.summary()
    return result


def test_brand_profile_to_format_spec():
    """测试Brand Profile转换为旧格式"""
    result = TestResult("brand_profile_to_format_spec")
    
    profile_path = Path(__file__).parent.parent / 'config' / 'brand_profile.json'
    if profile_path.exists():
        profile = load_brand_profile(str(profile_path))
        format_spec = brand_profile_to_format_spec(profile)
        
        # 检查转换结果
        result.assert_true('page' in format_spec, "有page字段")
        result.assert_true('heading1' in format_spec, "有heading1字段")
        result.assert_true('body' in format_spec, "有body字段")
        result.assert_true('figure_caption' in format_spec, "有figure_caption字段")
    else:
        result.assert_true(False, f"Brand Profile文件不存在: {profile_path}")
    
    result.summary()
    return result


# ============================================================================
# 集成测试 - generate_manual.py
# ============================================================================

def test_generate_manual_integration():
    """测试generate_manual.py的完整流程"""
    result = TestResult("generate_manual_integration")
    
    # 准备测试数据
    test_content = {
        "software_name": "测试软件",
        "version": "V1.0",
        "company": "测试公司",
        "cover": {
            "title": "测试软件V1.0",
            "subtitle": "用户使用手册",
            "company": "测试公司",
            "date": "2026年6月"
        },
        "sections": [
            {
                "heading": "一、系统概述",
                "content": "这是测试内容。",
                "subsections": [
                    {
                        "heading": "（一）功能介绍",
                        "content": "这是功能介绍内容。",
                        "images": []
                    }
                ]
            }
        ]
    }
    
    # 创建临时文件
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
        json.dump(test_content, f, ensure_ascii=False)
        content_path = f.name
    
    template_path = Path(__file__).parent.parent / 'templates' / 'manual_template_standard.docx'
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name
    
    try:
        # 导入generate_manual
        from generate_manual import ManualGenerator
        
        # 生成文档
        generator = ManualGenerator(str(template_path), str(Path(__file__).parent.parent / 'config' / 'brand_profile.json'))
        generator.generate(test_content, output_path)
        
        # 验证生成的文档
        doc = Document(output_path)
        
        # 检查文档不为空
        result.assert_true(len(doc.paragraphs) > 0, "文档不为空")
        
        # 检查封面标题
        cover_title = doc.paragraphs[0].text.strip()
        result.assert_equal(cover_title, "测试软件V1.0", "封面标题正确")
        
        # 检查章节标题
        has_heading1 = any(p.style.name == 'Heading 1' for p in doc.paragraphs)
        result.assert_true(has_heading1, "有Heading 1样式")
        
        has_heading2 = any(p.style.name == 'Heading 2' for p in doc.paragraphs)
        result.assert_true(has_heading2, "有Heading 2样式")
        
    finally:
        # 清理临时文件
        os.unlink(content_path)
        os.unlink(output_path)
    
    result.summary()
    return result


# ============================================================================
# 验收测试 - 文档符合Brand Profile规范
# ============================================================================

def test_document_conforms_to_brand_profile():
    """验证生成的文档符合Brand Profile规范"""
    result = TestResult("document_conforms_to_brand_profile")
    
    profile_path = Path(__file__).parent.parent / 'config' / 'brand_profile.json'
    doc_path = Path("/home/buqx/output_data/国农臻汇软著V2/output/test_brand_profile.docx")
    
    if not profile_path.exists():
        result.assert_true(False, f"Brand Profile不存在: {profile_path}")
        result.summary()
        return result
    
    if not doc_path.exists():
        result.assert_true(False, f"测试文档不存在: {doc_path}")
        result.summary()
        return result
    
    profile = load_brand_profile(str(profile_path))
    doc = Document(str(doc_path))
    
    # 1. 检查封面标题格式
    cover_style = profile['styles']['cover_title']
    for para in doc.paragraphs[:5]:
        text = para.text.strip()
        if text and '国农臻汇' in text:
            if para.runs:
                actual_size = round(para.runs[0].font.size.pt, 1)
                expected_size = cover_style['font']['size_pt']
                result.assert_equal(actual_size, expected_size, "封面标题字号")
                
                actual_bold = para.runs[0].font.bold
                expected_bold = cover_style['font']['bold']
                result.assert_equal(actual_bold, expected_bold, "封面标题加粗")
            break
    
    # 2. 检查正文格式
    body_style = profile['styles']['body']
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name
        if text and len(text) > 50 and style == 'Normal':
            actual_indent = round(para.paragraph_format.first_line_indent.cm, 2) if para.paragraph_format.first_line_indent else None
            expected_indent = body_style['paragraph']['first_line_indent_cm']
            result.assert_equal(actual_indent, expected_indent, "正文首行缩进")
            break
    
    # 3. 检查图名格式
    fig_style = profile['styles']['figure_caption']
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('图8'):
            if para.runs:
                actual_size = round(para.runs[0].font.size.pt, 1)
                expected_size = fig_style['font']['size_pt']
                result.assert_equal(actual_size, expected_size, "图名字号")
            
            actual_align = para.paragraph_format.alignment
            expected_align = {'center': 1, 'left': 0, 'right': 2, 'justify': 3}.get(fig_style['paragraph']['alignment'], 0)
            result.assert_equal(actual_align, expected_align, "图名对齐方式")
            break
    
    result.summary()
    return result


# ============================================================================
# 运行所有测试
# ============================================================================

def run_all_tests():
    """运行所有测试"""
    print("=" * 60)
    print("TDD测试套件 - 软著文档生成Skill")
    print("=" * 60)
    
    results = []
    
    # 单元测试
    print("\n--- 单元测试 ---")
    results.append(test_count_chars())
    results.append(test_extract_figure_refs())
    results.append(test_has_figure_caption())
    results.append(test_validate_content())
    results.append(test_brand_profile_loading())
    results.append(test_brand_profile_to_format_spec())
    
    # 集成测试
    print("\n--- 集成测试 ---")
    results.append(test_generate_manual_integration())
    
    # 验收测试
    print("\n--- 验收测试 ---")
    results.append(test_document_conforms_to_brand_profile())
    
    # 总结
    print("\n" + "=" * 60)
    total_passed = sum(r.passed for r in results)
    total_failed = sum(r.failed for r in results)
    total_tests = total_passed + total_failed
    
    if total_failed == 0:
        print(f"✅ 全部通过: {total_passed}/{total_tests}")
    else:
        print(f"❌ 有失败: {total_passed}/{total_tests} 通过, {total_failed} 失败")
        print("\n失败的测试:")
        for r in results:
            if r.failed > 0:
                print(f"   - {r.name}: {r.failed} 个失败")
                for error in r.errors:
                    print(f"     · {error}")
    
    return total_failed == 0


if __name__ == "__main__":
    success = run_all_tests()
    exit(0 if success else 1)
